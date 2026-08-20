"""Tests for text extraction and chunking (no OCR/LibreOffice)."""

from io import BytesIO

import pymupdf
import pytest
from docx import Document

from app.core.config import Settings
from app.services.chunking import chunk_text, chunk_text_structured
from app.services.extraction import ExtractionError, TextExtractor


def _settings(**overrides) -> Settings:
    base = {
        "database_url": "postgresql+asyncpg://u:p@localhost/db",
        "ocr_languages": "eng",
        "pdf_ocr_min_chars": 20,
        "chunk_size": 80,
        "chunk_overlap": 20,
        "libreoffice_path": "soffice",
    }
    base.update(overrides)
    return Settings(**base)  # type: ignore[arg-type]


def test_extract_docx() -> None:
    document = Document()
    document.add_paragraph("Первый абзац про уход за кожей.")
    document.add_paragraph("Второй абзац с рекомендациями.")
    buffer = BytesIO()
    document.save(buffer)

    blocks = TextExtractor(_settings()).extract("note.docx", buffer.getvalue())
    assert isinstance(blocks, list)
    assert blocks and isinstance(blocks[0], dict)
    text = "\n".join(str(b["content"]) for b in blocks)  # type: ignore[index]
    assert "Первый абзац" in text
    assert "Второй абзац" in text


def test_extract_docx_structured_headings_and_lists() -> None:
    document = Document()
    document.add_heading("Уход за кожей летом", level=1)
    document.add_heading(
        "Правила использования солнцезащитных средств",
        level=2,
    )
    for item in [
        "1. Наносите SPF за 20 минут до выхода.",
        "2. Обновляйте каждые 2 часа.",
        "3. Не забывайте зону ушей и шеи.",
        "4. Используйте водостойкие формулы у воды.",
        "5. Наносите достаточный слой.",
        "6. Повторяйте после купания.",
        "7. Храните средство в прохладе.",
    ]:
        document.add_paragraph(item)
    buffer = BytesIO()
    document.save(buffer)

    blocks = TextExtractor(_settings()).extract("summer.docx", buffer.getvalue())
    assert isinstance(blocks[0], dict)
    typed = blocks  # type: ignore[assignment]
    assert typed[0]["type"] == "heading"
    assert typed[0]["content"] == "Уход за кожей летом"
    assert typed[1]["type"] == "heading"
    assert "солнцезащитных" in typed[1]["content"]
    list_items = [b for b in typed if b["type"] == "list"]
    assert len(list_items) == 7

    chunks = chunk_text(
        typed,
        settings=_settings(chunk_size=1500, chunk_overlap=200),
    )
    sunscreen = [c for c in chunks if "солнцезащитных" in c]
    assert len(sunscreen) == 1
    assert "1. Наносите SPF" in sunscreen[0]
    assert "7. Храните средство" in sunscreen[0]


def test_dod_sunscreen_rules_single_chunk() -> None:
    """DoD: search hit content must contain heading + all 7 SPF rules."""
    document = Document()
    document.add_heading("Уход за кожей летом", level=1)
    document.add_paragraph("Вводный абзац про летний уход.")
    document.add_heading(
        "Правила использования солнцезащитных средств",
        level=2,
    )
    rules = [
        "1. Наносите SPF за 20 минут до выхода.",
        "2. Обновляйте каждые 2 часа.",
        "3. Не забывайте зону ушей и шеи.",
        "4. Используйте водостойкие формулы у воды.",
        "5. Наносите достаточный слой.",
        "6. Повторяйте после купания.",
        "7. Храните средство в прохладе.",
    ]
    for item in rules:
        document.add_paragraph(item)
    document.add_heading("Другая секция", level=2)
    document.add_paragraph("Не относится к SPF.")
    buffer = BytesIO()
    document.save(buffer)

    blocks = TextExtractor(_settings()).extract(
        "Уход за кожей летом.docx",
        buffer.getvalue(),
    )
    structured = chunk_text_structured(
        blocks,  # type: ignore[arg-type]
        settings=_settings(chunk_size=1500),
    )
    hit = next(
        c
        for c in structured
        if "правила" in c["content"].casefold()
        and "солнцезащитных" in c["content"].casefold()
    )
    content = hit["content"]
    assert "Правила использования солнцезащитных средств" in content
    for rule in rules:
        assert rule in content
    assert hit["metadata"]["section_title"] == (
        "Правила использования солнцезащитных средств"
    )
    assert hit["metadata"]["is_heading_only"] is False


def test_extract_pdf_text_layer() -> None:
    pdf = pymupdf.open()
    page = pdf.new_page()
    # Default PDF fonts are Latin-only; keep ASCII for reliable text-layer tests.
    page.insert_text((72, 72), "PDF text layer about peeling protocol")
    data = pdf.tobytes()
    pdf.close()

    blocks = TextExtractor(_settings()).extract("guide.pdf", data)
    assert any("peeling" in block.lower() for block in blocks)  # type: ignore[union-attr]


def test_extract_unsupported_extension() -> None:
    with pytest.raises(ExtractionError, match="Unsupported"):
        TextExtractor(_settings()).extract("notes.txt", b"hello")


def test_chunk_text_respects_size_and_overlap() -> None:
    settings = _settings(chunk_size=50, chunk_overlap=10)
    long_text = " ".join([f"word{i}" for i in range(40)])
    chunks = chunk_text([long_text], settings=settings)
    assert len(chunks) >= 2
    assert all(len(chunk) <= 100 for chunk in chunks)


def test_chunk_text_keeps_short_block() -> None:
    chunks = chunk_text(["короткий блок"], settings=_settings())
    assert chunks == ["короткий блок"]
